# -*- coding: utf-8 -*-
"""生成《高中数学教学能力体系分析报告》DOCX。"""
import os
import zipfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Twips

FONT = "Microsoft YaHei"
DARK = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x59, 0x59, 0x59)
LIGHT_GRAY = RGBColor(0x7F, 0x7F, 0x7F)
BODY_INK = RGBColor(0x33, 0x33, 0x33)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
TOTAL_DXA = 9360


def set_font(run, size=11, bold=False, color=BODY_INK, font=FONT, italic=False):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def style_font(style, size=11, bold=False, color=BODY_INK, font=FONT):
    style.font.name = font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    style_font(normal, 11, color=BODY_INK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    title = styles["Title"]
    style_font(title, 22, bold=True, color=DARK)
    tp = title.paragraph_format
    tp.space_before = Pt(0)
    tp.space_after = Pt(4)
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = styles["Subtitle"]
    style_font(subtitle, 13, color=RGBColor(0x44, 0x54, 0x6A))
    sp = subtitle.paragraph_format
    sp.space_before = Pt(0)
    sp.space_after = Pt(10)
    sp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h1 = styles["Heading 1"]
    style_font(h1, 16, bold=True, color=BLUE)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    style_font(h2, 13, bold=True, color=BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    style_font(h3, 12, bold=True, color=DARK_BLUE)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    bullet = styles.add_style("CxBullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = styles["Normal"]
    style_font(bullet, 11, color=BODY_INK)
    bp = bullet.paragraph_format
    bp.left_indent = Inches(0.5)
    bp.first_line_indent = Inches(-0.25)
    bp.space_after = Pt(8)
    bp.line_spacing = 1.167
    _add_style_numbering(bullet, 1)

    num = styles.add_style("CxNum", WD_STYLE_TYPE.PARAGRAPH)
    num.base_style = styles["Normal"]
    style_font(num, 11, color=BODY_INK)
    np = num.paragraph_format
    np.left_indent = Inches(0.5)
    np.first_line_indent = Inches(-0.25)
    np.space_after = Pt(8)
    np.line_spacing = 1.167
    _add_style_numbering(num, 2)

    cap = styles.add_style("CxTableCaption", WD_STYLE_TYPE.PARAGRAPH)
    cap.base_style = styles["Normal"]
    style_font(cap, 9.5, bold=True, color=GRAY)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)


def _add_style_numbering(style, num_id):
    ppr = style.element.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.insert_element_before(
        numpr, "w:pBdr", "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr"
    )


def add_para(doc, text, size=11, bold=False, color=BODY_INK, align="left", space_after=6):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="CxBullet")
        r = p.add_run(it)
        set_font(r, 11, color=BODY_INK)


def add_numbers(doc, items):
    for it in items:
        p = doc.add_paragraph(style="CxNum")
        r = p.add_run(it)
        set_font(r, 11, color=BODY_INK)


def add_h1(doc, text):
    doc.add_paragraph(text, style="Heading 1")


def add_h2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def add_h3(doc, text):
    doc.add_paragraph(text, style="Heading 3")


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "2E74B5")
    pbdr.append(left)
    ppr.insert_element_before(shd, "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr")
    ppr.insert_element_before(pbdr, "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr")
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(label + "　")
    set_font(r1, 11, True, RGBColor(0x1F, 0x3A, 0x5F))
    r2 = p.add_run(text)
    set_font(r2, 11, color=BODY_INK)


def add_table(doc, caption, headers, rows, widths, aligns=None):
    p = doc.add_paragraph(style="CxTableCaption")
    r = p.add_run(caption)
    set_font(r, 9.5, True, GRAY)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    _apply_table_geometry(table, widths)

    if aligns is None:
        aligns = ["left"] * len(headers)
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h, bold=True, size=10, align="center",
                 fill=HEADER_FILL, color=RGBColor(0x1F, 0x3A, 0x5F))
    _repeat_header(table.rows[0])
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell(table.rows[i + 1].cells[j], val, size=10, align=aligns[j])
    return table


def _apply_table_geometry(table, widths):
    assert sum(widths) == TOTAL_DXA, "列宽之和必须等于 9360 DXA"
    tbl = table._tbl
    old_tblpr = tbl.tblPr
    tblpr_xml = (
        '<w:tblPr %s>'
        '<w:tblW w:w="%d" w:type="dxa"/>'
        '<w:jc w:val="left"/>'
        '<w:tblInd w:w="120" w:type="dxa"/>'
        '<w:tblBorders>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
        '<w:tblLayout w:type="fixed"/>'
        '<w:tblCellMar>'
        '<w:top w:w="80" w:type="dxa"/>'
        '<w:left w:w="120" w:type="dxa"/>'
        '<w:bottom w:w="80" w:type="dxa"/>'
        '<w:right w:w="120" w:type="dxa"/>'
        '</w:tblCellMar>'
        '<w:tblLook w:val="04A0" w:firstRow="1" w:lastRow="0" w:firstColumn="1" '
        'w:lastColumn="0" w:noHBand="0" w:noVBand="1"/>'
        '</w:tblPr>'
    )
    new_tblpr = parse_xml(tblpr_xml % (nsdecls("w"), TOTAL_DXA))
    old_tblpr.addprevious(new_tblpr)
    old_tblpr.getparent().remove(old_tblpr)

    for j, w in enumerate(widths):
        table.columns[j].width = Twips(w)
        for row in table.rows:
            row.cells[j].width = Twips(w)


def _repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trpr.append(th)


def set_cell(cell, text, bold=False, size=10, align="left", fill=None, color=BODY_INK):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_font(r, size, bold, color)
    if fill:
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        tcpr.insert_element_before(
            shd, "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"
        )


def add_header_footer(doc):
    section = doc.sections[0]
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ppr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "BFBFBF")
    pbdr.append(bottom)
    ppr.insert_element_before(pbdr, "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr")
    r = hp.add_run("高中数学教学能力体系分析报告")
    set_font(r, 9, color=LIGHT_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = fp.add_run("第 ")
    set_font(r1, 9, color=LIGHT_GRAY)
    _add_field(fp, "PAGE")
    r2 = fp.add_run(" 页")
    set_font(r2, 9, color=LIGHT_GRAY)


def _add_field(paragraph, instr):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "7F7F7F")
    rpr.append(sz)
    rpr.append(color)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(rpr)
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


NUMBERING_XML = (
    '<w:numbering %s>'
    '<w:abstractNum w:abstractNumId="0">'
    '<w:multiLevelType w:val="singleLevel"/>'
    '<w:lvl w:ilvl="0">'
    '<w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="\u2022"/>'
    '<w:lvlJc w:val="left"/>'
    '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
    '<w:rPr><w:rFonts w:ascii="Microsoft YaHei" w:hAnsi="Microsoft YaHei" '
    'w:eastAsia="Microsoft YaHei"/><w:sz w:val="22"/></w:rPr>'
    '</w:lvl></w:abstractNum>'
    '<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
    '<w:abstractNum w:abstractNumId="1">'
    '<w:multiLevelType w:val="singleLevel"/>'
    '<w:lvl w:ilvl="0">'
    '<w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%%1."/>'
    '<w:lvlJc w:val="left"/>'
    '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
    '<w:rPr><w:rFonts w:ascii="Microsoft YaHei" w:hAnsi="Microsoft YaHei" '
    'w:eastAsia="Microsoft YaHei"/></w:rPr>'
    '</w:lvl></w:abstractNum>'
    '<w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num>'
    '</w:numbering>' % nsdecls("w")
)


def patch_numbering(docx_path):
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path) as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in zin.namelist():
                if name == "word/numbering.xml":
                    continue
                data = zin.read(name)
                if name == "[Content_Types].xml":
                    s = data.decode("utf-8")
                    if "numbering+xml" not in s:
                        s = s.replace(
                            "</Types>",
                            '<Override PartName="/word/numbering.xml" '
                            'ContentType="application/vnd.openxmlformats-officedocument.'
                            'wordprocessingml.numbering+xml"/></Types>',
                        )
                    data = s.encode("utf-8")
                elif name == "word/_rels/document.xml.rels":
                    s = data.decode("utf-8")
                    if "numbering.xml" not in s:
                        s = s.replace(
                            "</Relationships>",
                            '<Relationship Id="rIdCxNumbering" '
                            'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
                            'relationships/numbering" Target="numbering.xml"/></Relationships>',
                        )
                    data = s.encode("utf-8")
                zout.writestr(name, data)
            zout.writestr("word/numbering.xml", NUMBERING_XML.encode("utf-8"))
    os.replace(tmp, docx_path)


def build():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_header_footer(doc)

    # ---------------- 封面标题区 ----------------
    doc.add_paragraph("高中数学教学能力体系分析报告", style="Title")
    doc.add_paragraph(
        "以“快速提分”与“思维素养提升”为双目标的教师能力模型与实践路径",
        style="Subtitle",
    )
    meta = add_para(
        doc,
        "适用对象：高中数学教师、教研组长、培训机构教学团队　|　文档类型：能力分析报告　|　版本：V1.0　|　编制日期：2026年8月",
        size=9.5,
        color=LIGHT_GRAY,
        space_after=14,
    )
    ppr = meta._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E74B5")
    pbdr.append(bottom)
    ppr.insert_element_before(pbdr, "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr")

    # ---------------- 摘要 ----------------
    add_h1(doc, "摘要：核心结论与阅读指引")
    add_callout(
        doc,
        "核心结论",
        "高中数学教学要实现“快速提分”与“思维逻辑提升”的双重目标，教师需要的能力不是若干孤立技巧的叠加，"
        "而是一个由目标反推出来的五域结构：学科本体能力（教得对）、教学转化能力（教得懂）、思维训练能力（教得深）、"
        "应试指导能力（考得好）、教学管理能力（管得久）。五域能力共同作用，才能形成“诊断—归因—定标—方案—执行—评估”的提分闭环。",
    )
    add_para(
        doc,
        "阅读指引：本报告采用可追溯的逻辑链展开——第一部分从两个目标出发，推导出分数四因子模型，并据此映射出五域能力总纲；"
        "第二至第六部分逐域展开能力内涵与训练要点；第七部分给出落地路径（六步闭环与时间表）；"
        "第八部分对整套体系做逻辑自洽性检验；第九部分提供教师自评工具与优先级建议。"
        "建议先读第一部分建立整体框架，再按自身短板选择对应章节精读。",
    )

    # ---------------- 第一部分 ----------------
    add_h1(doc, "第一部分　逻辑起点：从目标反推能力")

    add_h2(doc, "1.1　双目标的统一性：提分与思维不是二选一")
    add_para(
        doc,
        "“快速提分”是短周期、显性、可测的目标；“思维逻辑提升”是长周期、隐性、可持续的目标。二者看似方向不同，"
        "实则互为条件，推导如下：",
    )
    add_numbers(
        doc,
        [
            "前提一：高考数学以“能力立意”命题——基础题考知识再现，中档题考方法迁移，压轴题考思维品质。因此思维水平直接决定分数上限。",
            "前提二：分数提升会带来正反馈（信心增强、投入度提高、练习质量上升），从而支撑思维训练长期持续。",
            "推理一：由前提一，只训练套路而不训练思维，分数会很快触到“天花板”，且面对新情境题无从下手。",
            "推理二：由前提二，只讲思维而不做应试验证，反馈周期长，学生容易失去动机，教学难以持续。",
            "结论：最优路径是“以思维训练为主线、以应试训练为验证”，用考试数据反哺思维诊断，让两个目标互相成就。",
        ],
    )

    add_h2(doc, "1.2　分数从何而来：四因子分解模型")
    add_para(
        doc,
        "为了把“提分”从口号变成可操作的工程，先回答一个基础问题：高考数学成绩由什么决定？"
        "本报告将其分解为四个因子，且论证它们构成完备集——学生失分只有这四个去向，不存在第五类：",
    )
    add_bullets(
        doc,
        [
            "知识掌握度（K）：概念、公式、定理是否理解并内化，能否在正确情境下被提取；",
            "思维运用度（T）：面对陌生或变形问题，能否调用恰当的思想方法完成迁移与推理；",
            "规范表达度（S）：步骤是否完整、书写是否清晰、关键结论是否写到采分点；",
            "临场稳定度（P）：时间分配、做题顺序、检查习惯与心态调节是否可靠。",
        ],
    )
    add_para(
        doc,
        "四因子中任一因子为短板，都会直接拉低总分。由此得到一个重要推论：教师的能力结构必须与四个因子一一对应，"
        "否则就会出现“学生明明懂了却考不好”或“刷了很多题却原地踏步”的错配。",
    )
    add_table(
        doc,
        "表1　分数四因子与学生侧表现、教师侧能力域的对应关系",
        ["影响因子", "学生侧典型表现", "教师侧能力域"],
        [
            ["知识掌握度 K", "概念不清、公式记错、定理不会用、用错条件", "学科本体能力"],
            ["思维运用度 T", "会做原题不会变式、遇新题无思路、只记住答案", "思维训练能力"],
            ["规范表达度 S", "会做但丢步骤分、跳步、书写混乱、答非所问", "教学转化能力（示范与反馈）"],
            ["临场稳定度 P", "时间不够、会而不对、难题打乱节奏、紧张失分", "应试指导能力"],
            ["（支撑全部因子）", "计划混乱、错题反复、缺乏反馈、目标模糊", "教学管理能力"],
        ],
        [1500, 5060, 2800],
        ["center", "left", "center"],
    )

    add_h2(doc, "1.3　五域能力总纲：目标—路径—能力同构")
    add_para(
        doc,
        "将四因子对应到教师侧，即得到五域能力模型。每一域都回答一个核心问题，且都能落到可观察的教学行为上：",
    )
    add_table(
        doc,
        "表2　五域能力总纲",
        ["能力域", "定位语", "核心回答的问题"],
        [
            ["学科本体能力", "教得对", "教的内容是否准确、系统、深刻？"],
            ["教学转化能力", "教得懂", "学生能否听懂、学会、记住、会用？"],
            ["思维训练能力", "教得深", "学生能否举一反三、迁移创新？"],
            ["应试指导能力", "考得好", "学生能否在考场上把会做的题全部做对？"],
            ["教学管理能力", "管得久", "训练系统能否持续运转、数据化迭代？"],
        ],
        [1700, 1200, 6460],
        ["center", "center", "left"],
    )
    add_para(
        doc,
        "五域关系（自洽性说明）：学科本体是“地基”（内容正确性），教学转化是“桥梁”（教与学的接口），"
        "思维训练是“核心”（决定分数上限），应试指导是“出口”（让思维与知识显性化为分数），"
        "教学管理是“保障”（让前四域持续运转）。五域之间存在清晰的递进与支撑关系，缺一不可，也不冗余。",
    )

    # ---------------- 第二部分 ----------------
    add_h1(doc, "第二部分　学科本体能力：教得对")
    add_para(
        doc,
        "学科本体能力是一切教学方法有效的前提。内容若不准、不系统、不深刻，再好的课堂设计都会把学生带到错误的方向。",
    )

    add_h2(doc, "2.1　知识体系完整且结构化")
    add_para(
        doc,
        "高中数学并非零散知识点，而是三条主线加一个活动板块：函数主线（函数概念与性质、基本初等函数、导数、数列）、"
        "几何与代数主线（立体几何、解析几何、平面向量、不等式）、概率与统计主线，以及数学建模与探究活动。"
        "教师能力要求是：能脱离课本默画出完整知识网络图；能说出每个章节在整个体系中的位置与前后关联；"
        "能指出同一个知识点在不同模块中的出场方式（例如不等式既是工具又是考查对象）。",
    )

    add_h2(doc, "2.2　概念本质理解：知其然更知其所以然")
    add_para(
        doc,
        "概念的本质理解决定教学深度。举例说明：",
    )
    add_bullets(
        doc,
        [
            "导数：本质是“瞬时变化率”，背后是极限思想。理解本质，才能讲清为什么导数可以判断单调性、求最值、刻画切线；只背求导公式，学生遇到概念题必错。",
            "函数单调性：本质是“对任意 x1＜x2，有 f(x1)＜f(x2)”的结构化不等式。证明过程的规范与否，就是逻辑链完整与否的直接体现。",
            "向量：本质是“既有大小又有方向”的运算对象与运算规则的扩展。理解其“数形双重身份”，才能打通几何与代数。",
        ],
    )

    add_h2(doc, "2.3　命题研究与考情把握")
    add_para(
        doc,
        "教师需要建立对高考试题的统计直觉：主干知识反复考查（函数与导数、几何与代数约占六成以上），"
        "基础、中档、压轴大致呈 4:4:2 的难度结构（经验估计，具体以当年试卷为准）；"
        "新高考呈现情境化、开放性、结构不良与跨章节综合等新动向。教师应能说出近三年本省市试卷的考点分布，"
        "并据此安排教学重心，而不是平均用力。",
    )

    add_h2(doc, "2.4　运算基本功")
    add_para(
        doc,
        "运算能力是高中数学的隐性门槛。教师需具备三个层次：算得快（常用结论、速算技巧）、算得准（草稿纸规划、检验意识）、"
        "算得巧（合理选择运算路径，如换元、整体代换、消元策略）。许多学生“会而不对”，根子在运算基本功，而非思维。",
    )

    add_h2(doc, "2.5　自检清单")
    add_bullets(
        doc,
        [
            "能否不看课本画出高中数学完整知识网络图并讲清主线关系？",
            "能否为每个核心概念（函数、导数、向量、概率）讲出“为什么这样定义”？",
            "能否说出近三年本省市高考的考点分布与难度结构？",
            "能否现场规范完成一道压轴题的完整解答并指出每一步的依据？",
        ],
    )

    # ---------------- 第三部分 ----------------
    add_h1(doc, "第三部分　教学转化能力：教得懂")
    add_para(
        doc,
        "学科知识是教师的“存货”，教学转化能力决定“存货”能否真正进入学生的认知结构。"
        "这一域是所有课堂行为的入口，也是归因正确与否的分水岭。",
    )

    add_h2(doc, "3.1　学情诊断四维模型：测、问、看、析")
    add_bullets(
        doc,
        [
            "测：用精准测试定位知识漏洞（章节诊断、小题定位，而不是整套卷子盲目刷）；",
            "问：追问学生的思路过程——“你当时为什么想到这个方法？”，暴露思维断点；",
            "看：观察做题习惯——草稿纸是否有序、步骤是否完整、时间如何分配；",
            "析：归因分类——不会做／会而不熟／熟而不准／准而不规范／规范而心态失衡。",
        ],
    )
    add_para(
        doc,
        "归因是逻辑自洽的第一道检验：归因错误，对策必然错误。例如把“知识性失分”当成“粗心”，"
        "把“方法性失分”当成“练得少”，都会导致无效训练。",
    )

    add_h2(doc, "3.2　讲解转化：把抽象讲具体")
    add_bullets(
        doc,
        [
            "类比法：数列类比函数、向量类比数轴上的有向线段，借旧知建构新知；",
            "图示法：先画图后推理，让数形结合成为课堂的默认动作；",
            "拆步法：把长题拆成标准动作链（读题—翻译—建模—求解—检验），降低认知负荷；",
            "反例法：用反例划定概念边界（如“单调区间不能并”用反例说明），加深理解。",
        ],
    )

    add_h2(doc, "3.3　变式与迁移：一题多变、多题归一")
    add_para(
        doc,
        "变式是连接“学会”与“会迁移”的桥梁。常见变式维度：条件变式（改参数范围）、结论变式（改设问角度）、"
        "情境变式（换生活背景）、结构变式（换函数形式）。变式必须指向明确的思维训练目标，"
        "否则就退化为题海。例如一道导数压轴题，可依次做三次变式后，让学生归纳“恒成立问题的通法框架”，"
        "实现从具体题到模型的收敛。",
    )

    add_h2(doc, "3.4　示范与规范：教师的板书就是标准答案的思维外显")
    add_para(
        doc,
        "学生的规范表达来自教师的示范。要求：先列思路再动笔，关键步骤不跳步，结论书写完整；"
        "课堂上至少完整示范一道中档题的“标准解答全过程”，让学生看到“想”与“写”之间的翻译规则。",
    )

    add_h2(doc, "3.5　分层适配：让每个学生都在最近发展区")
    add_bullets(
        doc,
        [
            "按目标分层：及格线、中档提升、冲高分三个层次，各有各的“保底动作”与“冲刺动作”；",
            "按内容分层：例题、作业、拓展题分级设计，同一课堂内让不同层次都有可得性；",
            "分层不是降低标准，而是让每层学生都够得着“跳一跳”才能达到的目标。",
        ],
    )

    # ---------------- 第四部分 ----------------
    add_h1(doc, "第四部分　思维训练能力：教得深")
    add_para(
        doc,
        "思维训练是五域中的核心：它直接决定分数上限，也直接回应“提升思维逻辑能力”的目标。"
        "关键是思维训练必须依托具体知识内容展开——脱离内容的“思维体操”是无效的，这保证了本部分与前两部分逻辑自洽。",
    )

    add_h2(doc, "4.1　思维目标：六大核心素养的考试化落地")
    add_table(
        doc,
        "表3　核心素养与考试载体、培养场景的对应",
        ["核心素养", "考试载体", "培养场景"],
        [
            ["数学抽象", "函数概念、新定义题", "概念发生式教学，追问“为什么这样定义”"],
            ["逻辑推理", "证明题、压轴题第二、三问", "问题链驱动、学生讲题说理"],
            ["数学建模", "应用题、情境题", "建模活动课、真实问题拆解"],
            ["直观想象", "立体几何、函数图像", "画图训练、数形结合"],
            ["数学运算", "解析几何、数列", "算理讲解、运算路径优化"],
            ["数据分析", "概率统计大题", "数据解读、误差与决策分析"],
        ],
        [1500, 2960, 4900],
        ["center", "left", "left"],
    )

    add_h2(doc, "4.2　八大思维方法：高中数学思想方法体系")
    add_table(
        doc,
        "表4　八大思维方法的典型场景与常见误区",
        ["思维方法", "典型场景", "常见误区"],
        [
            ["函数与方程", "恒成立、求范围、零点问题", "只代公式，不懂函数视角"],
            ["数形结合", "零点、不等式、最值", "图画不准，只形不数"],
            ["分类讨论", "含参问题、绝对值、分段", "漏类或重复讨论"],
            ["转化与化归", "陌生问题变熟悉问题", "盲目转化，不等价变形"],
            ["特殊与一般", "找规律、验证猜想", "以特殊代替一般"],
            ["有限与无限", "数列极限、无穷递缩", "用直觉代替严谨"],
            ["整体与局部", "换元、整体代换", "陷入局部纠缠"],
            ["或然与必然", "概率统计", "混淆频率与概率"],
        ],
        [1900, 3800, 3660],
        ["center", "left", "left"],
    )

    add_h2(doc, "4.3　五步思维训练法")
    add_numbers(
        doc,
        [
            "概念发生式：让学生经历概念的“诞生”，先追问“为什么这样定义”，再进入定义与性质。",
            "问题链驱动：用连续追问把一道题拆成“观察—猜想—验证—推广”，让推理过程外显。",
            "一题多解（发散）：比较不同解法的优劣、适用条件，训练选择能力。",
            "多题一解（收敛）：从一组题中提炼通法与模型，训练抽象概括能力。",
            "思维外显与复盘：学生讲题、写“思路说明书”，错题按“错因—正确思路—同类题”三段式整理。",
        ],
    )

    add_h2(doc, "4.4　教学示例：函数单调性证明的问题链")
    add_para(
        doc,
        "以一道“证明函数 f(x)=x+1/x 在 (0,1) 上单调递减”为例，展示问题链如何承载完整推理：",
    )
    add_bullets(
        doc,
        [
            "Q1 定义是什么？（复现定义，建立证明目标）",
            "Q2 为什么要“任取 x1＜x2”？（体会任意性对严谨性的意义）",
            "Q3 如何把 f(x1)-f(x2) 变成可判断符号的形式？（变形方向：通分、因式分解）",
            "Q4 这些变形技巧的本质是什么？（把抽象符号运算转化为可比较的量）",
            "Q5 这个套路还能迁移到哪里？（不等式证明、数列单调性、函数最值）",
        ],
    )
    add_para(
        doc,
        "一道小题承载的是“定义—变形—判断—迁移”的完整逻辑链。坚持此类训练，学生获得的不是一道题的答案，"
        "而是一种可以复制到所有推理场景的思维程序。",
    )

    # ---------------- 第五部分 ----------------
    add_h1(doc, "第五部分　应试指导能力：考得好")
    add_para(
        doc,
        "应试指导不是“投机取巧”，而是把已经形成的能力在考场上稳定兑现。它与前两部分的关系是："
        "没有知识基础与思维训练，应试技巧是无根之木；有了它们，应试指导就是最后一公里的转化器。",
    )

    add_h2(doc, "5.1　命题研究")
    add_bullets(
        doc,
        [
            "结构研究：选择、填空、解答题的考点分布与分值权重；",
            "规律研究：主干知识轮换考查，稳定中有创新，压轴题常考“函数与导数”“解析几何”两大模块；",
            "方向研究：新定义、情境化、开放性、结构不良、多选题等新题型的应对方式。",
        ],
    )

    add_h2(doc, "5.2　应试策略")
    add_bullets(
        doc,
        [
            "时间分配：小题限时训练，大题保底优先，难题最后处理；",
            "做题顺序：先易后难，遇到卡壳超过时限立即跳过，保持节奏；",
            "取舍策略：考前就定好“放弃底线”（如压轴题最后一问），把时间留给能拿的分。",
        ],
    )

    add_h2(doc, "5.3　规范答题与采分点")
    add_para(
        doc,
        "评分按采分点给分，教师的规范教学必须包含：设、列、解、答四要素完整；关键步骤不跳步；"
        "定义域、单位、结论书写到位；用答题卡限时模拟训练书写质量。",
    )

    add_h2(doc, "5.4　考后数据闭环：让每次考试都变成教学信号")
    add_table(
        doc,
        "表5　失分类型—成因—对策—跟踪指标矩阵",
        ["失分类型", "主要成因", "教学对策", "跟踪指标"],
        [
            ["知识性失分", "概念漏洞、公式误用", "回归课本、专题补漏", "复测正确率"],
            ["方法性失分", "思路断层、不会迁移", "变式训练、通法归纳", "同类题独立完成率"],
            ["运算性失分", "算理不清、草稿混乱", "算理讲解、草稿规划", "运算错误率"],
            ["规范性失分", "步骤不全、跳步", "采分点示范、面批", "步骤得分率"],
            ["心理性失分", "紧张、时间失衡", "模拟训练、心态策略", "用时波动幅度"],
        ],
        [1400, 2300, 3160, 2500],
        ["center", "left", "left", "center"],
    )

    add_h2(doc, "5.5　临场心理训练")
    add_para(
        doc,
        "将模拟考试常态化并弱化对单次结果的奖惩，训练“会做的题不丢分”这一第一目标；"
        "教授简单的深呼吸与注意力锚定技巧，让学生在卡壳时能快速回到当前题目，而不是被情绪带偏。",
    )

    # ---------------- 第六部分 ----------------
    add_h1(doc, "第六部分　教学管理能力：管得久")
    add_para(
        doc,
        "前四域解决“会不会教”的问题，管理能力解决“能不能持续教到位”的问题。没有管理闭环，"
        "任何优秀教学都只是随机事件，无法复制、无法迭代。",
    )

    add_h2(doc, "6.1　学生档案与数据化管理")
    add_para(
        doc,
        "一人一档：知识雷达图（各章节掌握度）、错题统计、成绩曲线、注意力与习惯观察记录。"
        "档案的价值不在记录本身，而在“每次教学决策都有数据依据”。",
    )

    add_h2(doc, "6.2　错题管理闭环")
    add_para(
        doc,
        "错题必须经过“记录—归类—重做—变式—周清/月清”五步，且每道错题都要标注错因与正确思路。"
        "只抄题不加工的错题本是无效劳动；把错题按知识、方法、规范、心态四类归因，才能驱动精准训练。",
    )

    add_h2(doc, "6.3　目标管理")
    add_para(
        doc,
        "学习目标遵循 SMART 原则：具体、可测、可达、相关、有时限。"
        "例如“四周内把函数单调性相关题型的正确率从 60% 提升到 85%”，"
        "就同时具备了可执行性与可检验性，也为后续归因提供了基准线。",
    )

    add_h2(doc, "6.4　家校协同")
    add_para(
        doc,
        "定期向家长反馈“进步点—问题点—配合动作”三段式信息，让家长做环境支持者而非施压者，"
        "避免家校要求互相矛盾抵消教学效果。",
    )

    add_h2(doc, "6.5　教师自我进化")
    add_bullets(
        doc,
        [
            "教研复盘：每两周一次考情与学情分析会，用数据说话；",
            "题库建设：题目按知识点、难度、思维方法三标签分类，形成可复用的训练资源；",
            "持续学习：跟进课程标准、高考评价体系与命题趋势，保持内容的前沿性。",
        ],
    )

    # ---------------- 第七部分 ----------------
    add_h1(doc, "第七部分　快速提分的六步闭环：从诊断到迭代")
    add_para(
        doc,
        "“快速提分”的秘密不是跳过过程，而是缩短反馈周期。六步闭环把教学变成可迭代的系统，"
        "第六步的输出回到第一步的输入，形成自我优化的飞轮：",
    )
    add_table(
        doc,
        "表6　六步闭环的步骤、动作、能力与产出",
        ["步骤", "关键动作", "对应能力", "可测产出", "建议周期"],
        [
            ["1 诊断", "测试＋访谈＋作业分析", "教学转化（诊断）", "学情报告、失分矩阵", "第 1 周"],
            ["2 归因", "区分知识／方法／规范／心态", "教学转化（归因）", "归因清单", "第 1 周"],
            ["3 定标", "设定 SMART 目标与里程碑", "教学管理", "目标书", "第 1 周"],
            ["4 方案", "制定个性化学习方案", "五域综合", "教学计划", "第 1—2 周"],
            ["5 执行", "课堂＋练习＋错题＋周测", "五域综合", "周测数据", "第 2—8 周"],
            ["6 评估", "数据分析、方案迭代", "教学管理", "月度报告", "每月"],
        ],
        [1200, 2700, 2000, 1960, 1500],
        ["center", "left", "left", "center", "center"],
    )

    add_h2(doc, "7.1　三个时间尺度的目标")
    add_bullets(
        doc,
        [
            "4 周：建立规范与习惯，补齐最致命漏洞，预期“会而不对”的问题明显减少；",
            "1 学期：中档题正确率显著提升，成绩进入稳定上升通道；",
            "1 年：思维品质发生质变，陌生情境迁移能力增强，压轴题实现突破。",
        ],
    )

    # ---------------- 第八部分 ----------------
    add_h1(doc, "第八部分　思维链自检：逻辑自洽性验证")
    add_para(
        doc,
        "整套体系的可靠性来自四个检验，任何一环不通过，说明推理链存在断裂：",
    )
    add_numbers(
        doc,
        [
            "目标可分解：五域能力与分数四因子一一映射，无遗漏、无冗余——每一个能力域都能指出它提升的是哪个因子。",
            "能力可培养：每一域都有具体的训练动作与载体（测试、变式、问题链、周测、错题闭环），不是空泛的口号。",
            "过程可验证：每一步都有可测指标（正确率、步骤得分率、用时、波动幅度），可以被数据检验。",
            "无循环论证：能力定义独立于“成绩提升”这一结果定义——先定义能力行为，再检验其对结果的影响，避免“因为他提分了，所以他能力强”的循环。",
        ],
    )
    add_h2(doc, "8.1　常见误区警示")
    add_bullets(
        doc,
        [
            "只刷题不反思：有量无质，训练的是机械反应而非思维；",
            "只讲技巧不讲原理：无根之木，题型一变即失效；",
            "只重结果不重归因：头痛医头，同类错误反复出现；",
            "只教知识不教思维：天花板过低，中档以上题目无法突破。",
        ],
    )

    # ---------------- 第九部分 ----------------
    add_h1(doc, "第九部分　教师能力自评与优先级建议")
    add_para(
        doc,
        "下表可作为教师自评工具：对每个关键行为按 1—5 分自评（1=基本不具备，5=成熟稳定），"
        "得分最低的两个维度即为下一阶段的提升重点。",
    )
    add_table(
        doc,
        "表7　五域能力自评表",
        ["能力域", "关键自评问题", "自评分（1—5）"],
        [
            ["学科本体", "能否脱稿画出完整知识网络图？能否讲清导数、向量等概念的来龙去脉？", ""],
            ["教学转化", "能否用测试加访谈准确归因学生失分？能否用变式让学生完成迁移？", ""],
            ["思维训练", "能否用问题链让一道题承载完整推理？能否让学生独立讲题说理？", ""],
            ["应试指导", "能否说出本省市近三年考点分布？能否用考后数据驱动教学调整？", ""],
            ["教学管理", "是否建立了学生档案与错题闭环？计划是否按 SMART 原则设定？", ""],
        ],
        [1700, 5860, 1800],
        ["center", "left", "center"],
    )
    add_h2(doc, "9.1　优先级建议")
    add_numbers(
        doc,
        [
            "第一步：诊断能力先行——没有诊断就没有针对性，一切训练都是盲目的；",
            "第二步：补齐学科本体——内容若有错误，教学方法越有效，危害越大；",
            "第三步：思维训练与应试指导并行——用应试数据验证思维训练效果；",
            "第四步：用管理闭环固化成果——把偶然的成功变成可复制的系统。",
        ],
    )

    # ---------------- 结语 ----------------
    add_h1(doc, "结语")
    add_callout(
        doc,
        "一句话总结",
        "好成绩是思维能力的副产品，好思维是教学设计的刻意结果。五域能力模型的意义，"
        "在于把“因材施教”从口号变成可诊断、可训练、可验证的工程——"
        "教师先把自己变成系统，学生才能在系统里稳定成长。",
    )

    # ---------------- 附录 ----------------
    add_h1(doc, "附录A　五域能力速查表")
    add_table(
        doc,
        "表8　五域能力速查",
        ["能力域", "子能力构成", "关键产出"],
        [
            ["学科本体", "知识体系、概念本质、命题研究、运算功底", "知识网络图、考点分布表"],
            ["教学转化", "学情诊断、讲解转化、变式迁移、规范示范、分层适配", "学情报告、变式题组"],
            ["思维训练", "素养落地、八大方法、五步训练法、问题链", "问题链教案、通法清单"],
            ["应试指导", "命题研究、应试策略、规范答题、数据闭环、心理训练", "失分矩阵、应试手册"],
            ["教学管理", "学生档案、错题闭环、目标管理、家校协同、自我进化", "月度报告、迭代计划"],
        ],
        [1600, 4260, 3500],
        ["center", "left", "left"],
    )

    add_h1(doc, "附录B　参考依据说明")
    add_para(
        doc,
        "本报告的能力框架以《普通高中数学课程标准》所确立的核心素养体系与《中国高考评价体系》的"
        "“一核四层四翼”为思想背景，结合一线教学中的经验规律（如难度结构 4:4:2 为经验估计，"
        "具体以当年试卷为准）整理而成。使用者应根据所在地区教材版本与高考方案做本地化调整。",
    )

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "高中数学教学能力体系分析报告.docx")
    doc.save(out)
    patch_numbering(out)
    print("已生成:", out)
    return out


if __name__ == "__main__":
    build()
