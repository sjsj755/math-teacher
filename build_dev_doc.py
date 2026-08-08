# -*- coding: utf-8 -*-
"""生成《数学学习提升App开发文档》DOCX。"""
import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_report import (
    BODY_INK,
    CALLOUT_FILL,
    DARK,
    GRAY,
    LIGHT_GRAY,
    add_bullets,
    add_callout,
    add_h1,
    add_h2,
    add_h3,
    add_numbers,
    add_para,
    add_table,
    configure_styles,
    patch_numbering,
    set_font,
    style_font,
)

CODE_FONT = "Consolas"


def add_code_style(doc):
    styles = doc.styles
    code = styles.add_style("CxCode", WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = styles["Normal"]
    code.font.name = CODE_FONT
    rpr = code.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), CODE_FONT)
    rfonts.set(qn("w:hAnsi"), CODE_FONT)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(0x1F, 0x2A, 0x36)
    pf = code.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Inches(0.12)
    pf.right_indent = Inches(0.12)
    return code


def add_code(doc, text, style=None):
    lines = text.split("\n")
    for i, line in enumerate(lines):
        p = doc.add_paragraph(style="CxCode")
        if style is not None:
            p = doc.add_paragraph(style="CxCode")
        ppr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F4F6F9")
        ppr.insert_element_before(shd, "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr")
        if i == 0:
            p.paragraph_format.space_before = Pt(6)
        if i == len(lines) - 1:
            p.paragraph_format.space_after = Pt(8)
        r = p.add_run(line if line else " ")
        set_font(r, 9, color=RGBColor(0x1F, 0x2A, 0x36), font=CODE_FONT)


def add_header_footer(doc, title):
    section = doc.sections[0]
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ppr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "BFBFBF")
    pbdr.append(bottom)
    ppr.insert_element_before(pbdr, "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr")
    r = hp.add_run(title)
    set_font(r, 9, color=LIGHT_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = fp.add_run("第 ")
    set_font(r1, 9, color=LIGHT_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r2 = OxmlElement("w:r")
    rpr2 = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "7F7F7F")
    rpr2.append(sz)
    rpr2.append(color)
    t = OxmlElement("w:t")
    t.text = "1"
    r2.append(rpr2)
    r2.append(t)
    fld.append(r2)
    fp._p.append(fld)
    r3 = fp.add_run(" 页")
    set_font(r3, 9, color=LIGHT_GRAY)


def build():
    doc = Document()
    configure_styles(doc)
    add_code_style(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_header_footer(doc, "数学学习提升 App 开发文档")

    # ---------------- 封面标题区 ----------------
    doc.add_paragraph("数学学习提升 App 开发文档", style="Title")
    doc.add_paragraph(
        "个人版 MVP（学生端＋家长 QQ 机器人日报）——架构、数据、接口与分步开发指南",
        style="Subtitle",
    )
    meta = add_para(
        doc,
        "文档版本：V1.3（MVP＋持续更新架构）　|　编制日期：2026 年 8 月　|　依据：《高中数学教学能力体系分析报告》及已确认产品决策　|　用途：后续“一步一步开发”的执行依据",
        size=9.5,
        color=LIGHT_GRAY,
        space_after=14,
    )
    ppr = meta._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E74B5")
    pbdr.append(bottom)
    ppr.insert_element_before(pbdr, "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr")

    # ---------------- 一、文档说明 ----------------
    add_h1(doc, "一、文档说明与使用方式")
    add_para(
        doc,
        "本文档是《数学学习提升 App》个人版 MVP 的完整开发依据，覆盖产品需求摘要、总体架构、数据设计、接口设计、"
        "学生端模块设计、核心流程、架构演进与持续更新、分步开发步骤、测试计划与迭代路线。后续开发将严格按照第九部分（开发步骤）的阶段顺序推进："
        "每完成一个阶段，先对照该阶段的“验收点”自检，通过后再进入下一阶段。",
    )
    add_callout(
        doc,
        "开发原则",
        "先跑通最小闭环，再逐步加功能。MVP 只实现“诊断→练习→错题→QQ 日报”，所有后续功能（目标管理、周测模考、"
        "AI 讲解与变式、云同步、教师端）都通过预留扩展点接入，不改动核心架构。",
    )
    add_para(
        doc,
        "修订记录：V1.3（2026-08-08）标记阶段 0、阶段 1 已完成，记录实际环境与构建配置，新增项目进度总览；"
        "V1.2（2026-08-07）架构升级为 DDD 有界上下文＋模块化单体，新增学习/内容/通知三域划分；"
        "V1.1（2026-08-07）新增“八、架构演进与持续更新”章节，接口统一升级为 /api/v1/，明确内容双轨、数据迁移、发布流程与测试基线。",
    )
    add_h2(doc, "1.1　术语表")
    add_table(
        doc,
        "表1　术语与含义",
        ["术语", "含义", "关联模块"],
        [
            ["四因子", "知识 K / 思维 T / 规范 S / 临场 P，成绩的四个决定因子", "诊断、报告"],
            ["六步闭环", "诊断→归因→定标→方案→执行→评估", "全局"],
            ["五域能力", "报告中的教师能力模型，转化为学生端功能依据", "全局"],
            ["MVP", "最小可行产品：诊断＋练习＋错题＋日报", "全局"],
            ["绑定码", "学生端生成的短码，家长用它完成 QQ 绑定", "家长通知"],
            ["openid", "家长在 QQ 机器人会话中的唯一标识", "服务端"],
            ["日报摘要", "每日上送的最小化学习数据（不含题目全文）", "家长通知"],
            ["扩展点", "预留的抽象接口，后续功能沿接口增量实现", "架构"],
            ["有界上下文", "DDD 中职责清晰的领域边界（学习/内容/通知三域）", "架构"],
            ["模块化单体", "单进程部署、模块边界清晰的应用形态，可逐步按域拆分", "架构"],
        ],
        [1600, 4960, 2800],
        ["center", "left", "center"],
    )
    add_h2(doc, "1.2　项目进度总览")
    add_callout(
        doc,
        "当前进度",
        "阶段 0（环境准备）✅ 已完成；阶段 1（工程骨架）✅ 已完成；阶段 2–6 待开发。下一阶段：阶段 2 数据层与题库。",
    )
    add_bullets(
        doc,
        [
            "✅ 已完成：Flutter 3.44.9（D:\\flutter）、Android SDK（D:\\Android，API 36）、iQOO Z5x 真机连接、QQ 官方机器人注册；",
            "✅ 已完成：math_up_app 工程骨架（依赖、中文主题、路由、6 个占位页），调试 APK 构建并真机运行验收通过；",
            "○ 待办：内网穿透注册（cpolar/花生壳）；阶段 2 数据层与题库 → 阶段 3 诊断报告 → 阶段 4 练习错题 → 阶段 5 通知服务 → 阶段 6 联调部署。",
        ],
    )

    # ---------------- 二、产品需求摘要 ----------------
    add_h1(doc, "二、产品需求摘要（MVP）")
    add_h2(doc, "2.1　目标与使用场景")
    add_bullets(
        doc,
        [
            "个人使用：一名高中生（学生端 App）＋一位家长（官方 QQ 机器人接收日报），无多租户、无商业化、无管理后台。",
            "学生目标：通过“诊断→练习→错题”闭环快速提分并记录思维薄弱点；",
            "家长目标：每天及时了解孩子完成情况，无需安装第二个 App。",
            "内容基准：四川地区、新高考 II 卷（3+1+2）、人教 A 版 2020 年 5 月第一版教材，覆盖高一至高三全学段。",
        ],
    )
    add_h2(doc, "2.2　MVP 功能清单")
    add_table(
        doc,
        "表2　功能范围：MVP 与后续版本",
        ["模块", "功能点", "版本"],
        [
            ["首次诊断", "选年级→自适应诊断（15 题）→四因子得分", "MVP"],
            ["学情报告", "四因子雷达图、薄弱知识点、归因清单", "MVP"],
            ["练习", "按薄弱知识点推荐、计时、即时解析", "MVP"],
            ["错题本", "自动收录、三类归因、重做与掌握状态", "MVP"],
            ["QQ 日报", "绑定、每日摘要推送、主动拉取、解绑", "MVP"],
            ["目标管理", "SMART 目标、4 周/学期/年里程碑", "V1.1"],
            ["周测模考", "按 II 卷结构组卷、限时、失分矩阵", "V1.2"],
            ["AI 讲解变式", "AI 生成讲解与变式题（人工审核）", "V1.3"],
            ["云同步/教师端", "账号体系、班级接口、日报历史页", "V1.4"],
        ],
        [1500, 5060, 2800],
        ["center", "left", "center"],
    )
    add_h2(doc, "2.3　非功能要求")
    add_bullets(
        doc,
        [
            "离线可用：无网时本地完成练习与错题记录，恢复网络后补传日报摘要；",
            "数据最小化：服务端只接收日报摘要，不接收题目全文与个人敏感信息；",
            "单用户：学生端无需登录，数据存本地 SQLite；",
            "性能：App 冷启动 < 2 秒，题目加载 < 1 秒，支持 60 帧滚动。",
        ],
    )

    # ---------------- 三、总体架构 ----------------
    add_h1(doc, "三、总体架构")
    add_para(
        doc,
        "架构形态：采用 DDD 领域驱动设计组织代码——先按有界上下文划分领域（学习域/内容域/通知域，详见 8.8），"
        "各域内分层建模；整体保持模块化单体（单进程部署），现阶段不拆微服务，但域边界与接口为未来按域拆分预留。",
    )
    add_h2(doc, "3.1　系统组成")
    add_para(doc, "系统由三个部分组成，数据流向如下：")
    add_code(
        doc,
        "学生端 App（Flutter，本地 SQLite）\n"
        "        │  ① 绑定码/状态查询  ② 日报摘要上送\n"
        "        ▼\n"
        "通知服务（FastAPI + SQLite + APScheduler，部署于个人设备）\n"
        "        │  ③ 绑定指令回调  ④ 日报消息推送（官方 QQ 机器人 API）\n"
        "        ▼\n"
        "家长 QQ（与官方机器人会话，接收日报）",
    )
    add_para(
        doc,
        "关键边界：机器人密钥与家长 openid 只保存在服务端，学生端 App 不与 QQ 平台直连；"
        "学生端与服务端之间只传输绑定信息与日报摘要。",
    )
    add_h2(doc, "3.2　技术选型")
    add_table(
        doc,
        "表3　技术选型与理由",
        ["技术", "用途", "选择理由"],
        [
            ["Flutter", "学生端跨平台 App", "一套代码覆盖 iOS/Android，数学公式与图表渲染生态成熟"],
            ["flutter_math_fork", "公式渲染", "LaTeX 风格公式显示，适合数学题干与解析"],
            ["fl_chart", "雷达图/趋势图", "报告页图表轻量实现"],
            ["sqflite/drift", "本地数据库", "SQLite 稳定可靠，离线优先"],
            ["Python FastAPI", "通知服务", "轻量、易扩展，AI 集成生态好"],
            ["APScheduler", "定时推送", "每日 21:30 触发日报任务"],
            ["官方 QQ 机器人 API", "家长通知", "合规通道，个人使用配额内可满足日推 1 条"],
            ["SQLite（服务端）", "绑定与日志", "单用户场景无需数据库服务"],
            ["Git＋semver 版本规范", "版本管理与变更记录", "内容、功能双轨持续更新的基线"],
            ["内网穿透（cpolar/花生壳）", "QQ Webhook 公网回调", "家用电脑部署时保证 QQ 平台可达"],
            ["DDD 领域驱动设计", "架构组织方式", "领域概念清晰、模块边界稳定，为按域拆分预留"],
        ],
        [1900, 3560, 3900],
        ["center", "left", "left"],
    )
    add_h2(doc, "3.3　分层架构")
    add_para(doc, "学生端采用 feature 优先的四层结构，服务端采用四层结构，模块之间只通过领域实体与接口交互：")
    add_code(
        doc,
        "学生端：presentation（页面/组件）→ application（用例）→ domain（实体/仓储接口）→ infrastructure（SQLite/题库 JSON）\n"
        "服务端：api（路由）→ application（用例/编排）→ domain（领域实体与规则）→ infrastructure（通道/仓储实现）",
    )
    add_h2(doc, "3.4　扩展点（后续增量功能接入位）")
    add_para(doc, "以下接口在 MVP 提供规则实现，后续版本直接替换实现类即可，不修改上层调用：")
    add_code(
        doc,
        "// 学生端（Dart）扩展点\n"
        "abstract class QuestionRepository { Future<List<Question>> byKnowledge(String code); }\n"
        "abstract class DiagnosticEngine { Future<DiagnosisResult> run(List<AnswerRecord> records); }\n"
        "abstract class Recommender { Future<List<Question>> recommend(DiagnosisResult d); }\n"
        "abstract class SyncService { Future<void> uploadDailyDigest(DailyDigest d); }\n"
        "abstract class ContentProvider { Future<ContentIndex> loadContent(); } // 内容源：本地 assets → 远程内容包\n"
        "\n"
        "# 服务端（Python）扩展点\n"
        "class NotificationChannel(ABC):\n"
        "    async def send(self, message: DigestMessage) -> bool: ...\n"
        "class QQBotChannel(NotificationChannel): ...   # MVP\n"
        "class EmailChannel(NotificationChannel): ...   # 预案",
    )
    add_para(
        doc,
        "内容与代码分离：题库与章节树作为独立内容包（content/），统一由 ContentProvider 提供；"
        "MVP 读本地 assets，未来切换远程内容包时只需替换该接口实现。",
    )
    add_para(
        doc,
        "扩展点与有界上下文对应：ContentProvider 属内容域，DiagnosticEngine/Recommender 属学习域，"
        "NotificationChannel 属通知域，SyncService 属共享支撑；各域通过防腐层接口交互，不直接依赖实现。",
    )
    add_h2(doc, "3.5　目录结构")
    add_para(
        doc,
        "仓库采用 monorepo 布局：app/（Flutter 学生端）、server/（通知服务）、content/（题库与章节树，独立版本）、"
        "docs/（文档与 CHANGELOG）。各端内部结构如下：",
    )
    add_code(
        doc,
        "math_up_app/                     # Flutter 学生端（DDD-lite）\n"
        "├─ lib/\n"
        "│  ├─ main.dart\n"
        "│  ├─ app.dart                  # 路由与主题\n"
        "│  ├─ core/\n"
        "│  │  ├─ db/database.dart        # infrastructure：SQLite 初始化\n"
        "│  │  ├─ domain/                 # 领域层\n"
        "│  │  │  ├─ models/              # 实体与值对象\n"
        "│  │  │  └─ repositories/        # 仓储接口\n"
        "│  │  ├─ infrastructure/         # SQLite/题库 JSON 实现\n"
        "│  │  └─ application/            # 用例（诊断/推荐/摘要）\n"
        "│  └─ features/\n"
        "│     ├─ onboarding/             # 首次启动选年级\n"
        "│     ├─ diagnosis/              # 诊断与报告\n"
        "│     ├─ practice/               # 练习\n"
        "│     ├─ errorbook/              # 错题本\n"
        "│     └─ settings/               # 家长通知设置\n"
        "├─ assets/data/chapters.json     # 章节树\n"
        "└─ assets/data/questions.json    # 题库\n"
        "\n"
        "notify_server/                   # FastAPI 通知服务（DDD 分层）\n"
        "├─ app/\n"
        "│  ├─ main.py                    # 入口\n"
        "│  ├─ config.py                  # QQ 凭据/推送时间\n"
        "│  ├─ schemas.py                 # 请求响应模型\n"
        "│  ├─ api/ bind.py digest.py qq_callback.py     # 路由层\n"
        "│  ├─ application/ digester.py scheduler.py      # 用例/编排\n"
        "│  ├─ domain/ entities.py                        # 领域实体与规则\n"
        "│  └─ infrastructure/\n"
        "│     ├─ channels/ qq_bot.py email.py            # 通道实现\n"
        "│     └─ repository.py                           # 数据仓储实现\n"
        "├─ requirements.txt\n"
        "└─ README.md",
    )

    # ---------------- 四、数据设计 ----------------
    add_h1(doc, "四、数据设计")
    add_h2(doc, "4.1　章节树与知识点编码（人教 A 版 2020 年 5 月第一版）")
    add_para(
        doc,
        "知识点编码规则：册代码-章-节，如 A1-3-2 表示必修第一册第 3 章第 2 节。"
        "题目、诊断、错题均使用该编码，保证全链路可追溯。",
    )
    add_bullets(
        doc,
        [
            "年级—章节映射（诊断与推荐按此取题）：高一＝必修第一、二册；高二＝选择性必修第一、二、三册；高三＝全部章节＋综合复习组。",
            "章节树文件 chapters.json 与题库 questions.json 必须使用同一套编码，导入时做一致性校验，杜绝悬空知识点。",
        ],
    )
    add_table(
        doc,
        "表4　五册章节树与编码前缀",
        ["册", "章节", "核心知识点", "编码前缀"],
        [
            ["必修第一册", "第1章 集合与常用逻辑用语", "集合运算、充分必要条件、量词", "A1-1"],
            ["必修第一册", "第2章 一元二次函数、方程和不等式", "不等式性质、基本不等式", "A1-2"],
            ["必修第一册", "第3章 函数的概念与性质", "函数概念、单调性、奇偶性", "A1-3"],
            ["必修第一册", "第4章 指数函数与对数函数", "指数对数运算与图像性质", "A1-4"],
            ["必修第一册", "第5章 三角函数", "任意角、三角公式、图像性质", "A1-5"],
            ["必修第二册", "第6章 平面向量及其应用", "向量运算、解三角形", "A2-6"],
            ["必修第二册", "第7章 复数", "复数四则运算、几何意义", "A2-7"],
            ["必修第二册", "第8章 立体几何初步", "空间几何体、点线面位置关系", "A2-8"],
            ["必修第二册", "第9章 统计", "抽样、频率分布、百分位数", "A2-9"],
            ["必修第二册", "第10章 概率", "古典概型、事件关系", "A2-10"],
            ["选择性必修第一册", "第1章 空间向量与立体几何", "空间向量、夹角与距离", "B1-1"],
            ["选择性必修第一册", "第2章 直线和圆的方程", "直线方程、圆、位置关系", "B1-2"],
            ["选择性必修第一册", "第3章 圆锥曲线的方程", "椭圆、双曲线、抛物线", "B1-3"],
            ["选择性必修第二册", "第4章 数列", "等差等比、数列求和", "B2-4"],
            ["选择性必修第二册", "第5章 一元函数的导数及其应用", "导数概念、单调性与极值", "B2-5"],
            ["选择性必修第三册", "第6章 计数原理", "排列组合、二项式定理", "B3-6"],
            ["选择性必修第三册", "第7章 随机变量及其分布", "分布列、期望与方差", "B3-7"],
            ["选择性必修第三册", "第8章 成对数据的统计分析", "回归分析、独立性检验", "B3-8"],
        ],
        [1900, 2960, 3100, 1400],
        ["center", "left", "left", "center"],
    )
    add_h2(doc, "4.2　题库 JSON 格式")
    add_para(
        doc,
        "题库文件 assets/data/questions.json 的题目结构如下（难度 1–5，1 最易；lose_type 取值 knowledge/method/calculation/standard/psychology）：",
    )
    add_code(
        doc,
        '{\n'
        '  "id": "A1-3-2-001",\n'
        '  "chapter": "A1-3",\n'
        '  "knowledge_point": "函数单调性的定义与判断",\n'
        '  "type": "choice",\n'
        '  "difficulty": 2,\n'
        '  "thinking_method": "逻辑推理",\n'
        '  "lose_type": "knowledge",\n'
        '  "stem": "设 f(x) 在区间 I 上单调递减，若 x1 < x2，则必有……",\n'
        '  "options": ["f(x1) > f(x2)", "f(x1) < f(x2)", "f(x1) = f(x2)", "无法确定"],\n'
        '  "answer": "A",\n'
        '  "explain": "由单调递减定义，x1 < x2 时 f(x1) > f(x2)。",\n'
        '  "variant_group": "A1-3-2"\n'
        "}",
    )
    add_h2(doc, "4.3　学生端 SQLite 表概览")
    add_table(
        doc,
        "表5　学生端数据表",
        ["表", "主要字段", "用途"],
        [
            ["question", "id/chapter/knowledge_point/type/difficulty/tags/answer/explain", "内置题库（导入自 JSON）"],
            ["diagnosis", "id/date/k_score/t_score/s_score/p_score/weak_points", "历次诊断结果"],
            ["answer_record", "id/question_id/result/seconds/date", "每题作答与用时"],
            ["error_book", "id/question_id/lose_type/status/redo_count", "错题与重做状态"],
            ["app_config", "key/value", "年级、推送时间、日报开关、绑定状态"],
            ["digest_queue", "id/date/payload/synced", "离线时缓存待上送的摘要"],
        ],
        [1800, 4160, 3400],
        ["center", "left", "left"],
    )
    add_h2(doc, "4.4　服务端表概览")
    add_table(
        doc,
        "表6　服务端数据表",
        ["表", "主要字段", "用途"],
        [
            ["binding", "device_id/bind_code/parent_openid/status/created_at", "学生设备与家长 openid 的绑定关系"],
            ["digest_log", "id/device_id/date/payload/created_at", "每日摘要记录"],
            ["push_log", "id/digest_id/channel/status/error/pushed_at", "推送结果与失败重试记录"],
        ],
        [1800, 4160, 3400],
        ["center", "left", "left"],
    )
    add_h2(doc, "4.5　四因子与归因规则（MVP 规则引擎）")
    add_numbers(
        doc,
        [
            "因子得分：诊断 15 题按因子分组（K 6 题、T 5 题、S 2 题、P 2 题），因子得分 = 该组题目正确率；S、P 组含自评题（步骤完整性、限时完成度）。",
            "薄弱知识点：某知识点题目正确率 < 60% 且出现 ≥ 2 题，则标记为薄弱点，按正确率升序排列。",
            "归因清单：错题按 lose_type 统计，前两名作为主归因；MVP 自动归因知识/方法/运算三类，规范/心理类通过作答后的自评选项采集。",
        ],
    )
    add_h2(doc, "4.6　数据迁移与内容版本")
    add_bullets(
        doc,
        [
            "schema_version：app_config 中固化数据库结构版本号，App 启动时与代码期望版本比对，低版本按序执行迁移脚本；",
            "content_version：content/ 内容包携带版本号，App 升级后自动检测并重导内容（仅覆盖题目数据，不影响学习记录）；",
            "升级必须采用“同签名覆盖安装”，禁止卸载重装，以保证 SQLite 学习数据连续。",
        ],
    )
    add_code(
        doc,
        "math_up_app/lib/core/db/migrations/\n"
        "├─ 001_init.sql          # 建表（表5）\n"
        "├─ 002_add_xxx.sql       # 后续结构变更（同步升 schema_version）\n"
        "└─ migrate.dart          # 按序执行，事务内完成",
    )

    # ---------------- 五、接口设计 ----------------
    add_h1(doc, "五、接口设计（通知服务）")
    add_h2(doc, "5.1　通用约定")
    add_bullets(
        doc,
        [
            "Base URL：部署地址＋版本前缀（如 http://192.168.1.10:8000/api/v1），学生端在设置页配置；",
            "数据格式：JSON；成功返回 200，参数错误 400，未找到绑定 404，冲突 409，服务错误 500；",
            "业务错误码：BIND_CODE_INVALID（绑定码无效）、ALREADY_BOUND（已绑定）、DEVICE_MISMATCH（设备不匹配）、NOT_BOUND（未绑定）。",
            "接口版本化：所有接口统一 /api/v1/ 前缀；新增字段必须向后兼容，破坏性变更必须升级 v2 并保留旧版过渡。",
        ],
    )
    add_h2(doc, "5.2　接口清单")
    add_table(
        doc,
        "表7　API 一览",
        ["方法", "路径", "用途", "调用方"],
        [
            ["POST", "/api/v1/bind", "学生端确认绑定码关联", "学生端"],
            ["GET", "/api/v1/status", "查询绑定与推送状态", "学生端"],
            ["POST", "/api/v1/daily-digest", "上送日报摘要", "学生端"],
            ["POST", "/api/v1/qq/callback", "接收家长 QQ 指令（绑定/日报/解绑）", "QQ 平台"],
            ["内部", "APScheduler 定时任务", "每日推送日报", "服务端"],
        ],
        [700, 2200, 3560, 2900],
        ["center", "left", "left", "center"],
    )
    add_h2(doc, "5.3　接口详述")
    add_h3(doc, "5.3.1　POST /api/v1/bind")
    add_code(
        doc,
        '请求：{"bind_code": "MATH-8F3K", "device_id": "uuid-of-device"}\n'
        '响应：{"status": "bound", "parent_nick": "家长昵称"}\n'
        "错误：404 BIND_CODE_INVALID / 409 ALREADY_BOUND",
    )
    add_h3(doc, "5.3.2　GET /api/v1/status")
    add_code(
        doc,
        '请求：GET /api/v1/status?device_id=uuid-of-device\n'
        '响应：{"bound": true, "last_push_at": "2026-08-07 21:30", '
        '"daily_enabled": true, "push_time": "21:30"}',
    )
    add_h3(doc, "5.3.3　POST /api/v1/daily-digest")
    add_code(
        doc,
        '请求：{"device_id": "uuid", "date": "2026-08-07", "practice_count": 10,\n'
        '       "correct_count": 8, "error_count": 2, "minutes": 45,\n'
        '       "weak_points": ["A1-3-2"], "streak_days": 3}\n'
        "响应：200 {\"accepted\": true}",
    )
    add_h3(doc, "5.3.4　POST /api/v1/qq/callback（QQ 平台 Webhook）")
    add_code(
        doc,
        '事件：C2C 消息，家长发送文本指令\n'
        '  "绑定 MATH-8F3K"  → 记录 openid 与绑定码关联\n'
        '  "日报"           → 立即推送当日摘要\n'
        '  "解绑"           → 解除绑定\n'
        "响应：200 确认收到",
    )
    add_h2(doc, "5.4　QQ 机器人接入说明")
    add_numbers(
        doc,
        [
            "在 QQ 开放平台注册开发者账号并创建机器人，获取 AppID、AppSecret 与 Token；",
            "配置 Webhook 回调地址指向 /api/v1/qq/callback（需公网可达或内网穿透）；",
            "家长在 QQ 中搜索并添加机器人，主动发送第一条消息（建立会话授权）；",
            "服务端使用官方消息发送 API 推送日报；注意平台配额：主动推送条数与能力受平台策略约束；",
            "预案：若定时推送不可用，降级为家长发“日报”指令主动拉取；邮件通道作为二线预案（代码位预留）。",
        ],
    )
    add_h3(doc, "5.4.1　日报消息模板")
    add_code(
        doc,
        "【数学学习日报】8月7日\n"
        "今日完成 10 题 · 正确率 80%\n"
        "错题 2 道：函数单调性、指数运算\n"
        "学习时长 45 分钟 · 连续打卡 3 天\n"
        "明日建议：先重做错题本中的 2 道题，再完成推荐练习。",
    )

    # ---------------- 六、学生端模块设计 ----------------
    add_h1(doc, "六、学生端模块设计")
    add_h2(doc, "6.1　页面地图")
    add_table(
        doc,
        "表8　页面与路由",
        ["页面", "路由", "说明"],
        [
            ["年级选择", "/onboarding", "首次启动，选择高一/高二/高三"],
            ["诊断页", "/diagnosis", "15 题顺序作答，含自评题"],
            ["报告页", "/report", "四因子雷达图、薄弱点、归因清单"],
            ["练习页", "/practice", "按推荐出题，计时与即时解析"],
            ["错题本", "/errorbook", "状态列表、重做、周清入口"],
            ["设置页", "/settings", "绑定码、推送时间、日报开关、解绑"],
        ],
        [1500, 3100, 4760],
        ["center", "left", "left"],
    )
    add_h2(doc, "6.2　诊断模块")
    add_para(
        doc,
        "诊断题配比固定为 K6 / T5 / S2 / P2。K 组按当前年级章节抽基础题；T 组抽带思维方法标签的中档题；"
        "S 组为“步骤完整性”自评题；P 组为限时完成题与心态自评。提交后调用 DiagnosticEngine 规则实现计算四因子得分。",
    )
    add_table(
        doc,
        "表9　诊断题配比",
        ["因子", "题数", "说明"],
        [
            ["知识 K", "6", "按年级章节抽基础题，测概念与公式"],
            ["思维 T", "5", "带思维方法标签的中档题，测迁移"],
            ["规范 S", "2", "自评：解答题步骤是否完整"],
            ["临场 P", "2", "限时题＋心态自评"],
        ],
        [1300, 1800, 6260],
        ["center", "center", "left"],
    )
    add_h2(doc, "6.3　练习模块（推荐规则）")
    add_numbers(
        doc,
        [
            "选题：按薄弱知识点正确率升序取前 3 个知识点，各知识点内按难度比例出题；",
            "难度比例：基础 60%（难度 1–2）、中档 30%（难度 3）、压轴 10%（难度 4–5）；",
            "单次练习 10 题；限时：选择题 2 分钟、填空题 3 分钟、解答题 10 分钟；",
            "提交后即时显示答案、解析与采分点；错误题自动写入错题本。",
        ],
    )
    add_h2(doc, "6.4　错题模块（状态机）")
    add_table(
        doc,
        "表10　错题状态流转",
        ["状态", "触发条件", "动作"],
        [
            ["待重做", "练习答错自动收录", "显示错题与解析"],
            ["已重做", "重做回答正确", "进入观察期（7 天后复测）"],
            ["已掌握", "观察期复测正确", "移出每日推荐，保留档案"],
            ["待重做（升级）", "重做或复测再错", "redo_count+1，重新推荐同类题"],
        ],
        [1800, 3760, 3800],
        ["center", "left", "left"],
    )
    add_h2(doc, "6.5　报告模块与设置模块")
    add_bullets(
        doc,
        [
            "报告：四因子雷达图（fl_chart）、薄弱知识点列表、近 7 天正确率趋势；",
            "设置：展示绑定码（生成后 24 小时有效）、推送时间（默认 21:30）、日报开关、解绑按钮；",
            "绑定状态轮询：设置页每 30 秒查询 /api/v1/status，家长完成 QQ 绑定后自动显示“已绑定”。",
        ],
    )

    # ---------------- 七、核心流程 ----------------
    add_h1(doc, "七、核心流程")
    add_h2(doc, "7.1　首次使用流程")
    add_numbers(
        doc,
        [
            "启动 App → 选择年级 → 进入诊断页；",
            "完成 15 题诊断（含自评）→ 生成四因子得分与薄弱点；",
            "展示报告页 → 自动跳转练习页，按薄弱点出 10 题；",
            "设置页生成绑定码，家长在 QQ 发送“绑定 绑定码”完成关联。",
        ],
    )
    add_h2(doc, "7.2　每日使用流程")
    add_numbers(
        doc,
        [
            "学生完成推荐练习与错题重做（可离线）；",
            "App 在本地生成日报摘要并写入 digest_queue；",
            "有网时自动上送 /api/v1/daily-digest；",
            "服务端 21:30 定时推送日报给家长 QQ；推送失败自动重试 3 次，仍失败次日合并提醒。",
        ],
    )

    # ---------------- 八、架构演进与持续更新 ----------------
    add_h1(doc, "八、架构演进与持续更新")
    add_para(
        doc,
        "本章为长期持续更新确立的架构机制：任何一次“改内容”或“改功能”都按本章流程发布，"
        "不破坏既有数据与接口，保证项目可持续演进。",
    )
    add_h2(doc, "8.1　版本管理规范")
    add_table(
        doc,
        "表11　三套版本号对照",
        ["对象", "版本号示例", "更新时机"],
        [
            ["App（学生端）", "v1.0.1", "代码或内容随包发布时递增（semver）"],
            ["内容包 content/", "content 12", "题库/章节树/解析变更时递增"],
            ["服务端 API", "/api/v1", "破坏性变更时升 v2"],
        ],
        [1900, 2800, 4660],
        ["center", "center", "left"],
    )
    add_bullets(
        doc,
        [
            "每版发布前更新 CHANGELOG（新增/修复/内容变更三栏）与版本号；",
            "升级必须“同签名覆盖安装”APK，禁止卸载重装，防止本地学习数据丢失。",
        ],
    )
    add_h2(doc, "8.2　内容与代码分离（双轨核心）")
    add_bullets(
        doc,
        [
            "content/ 目录独立存放 chapters.json、questions.json、content_index.json（含 content_version）；",
            "内容更新只动 content/ 与版本号，不改业务代码；构建时由 ContentProvider 统一加载；",
            "发布前运行内容一致性校验脚本：编码有效、题目 id 唯一、答案合法、无悬空知识点。",
        ],
    )
    add_h2(doc, "8.3　数据迁移机制")
    add_bullets(
        doc,
        [
            "SQLite schema_version 与 migrations/ 脚本目录配套，App 启动时自动按序迁移并保留数据；",
            "每次结构变更必须新增迁移脚本并升 schema_version，禁止直接改旧脚本。",
        ],
    )
    add_h2(doc, "8.4　API 版本化与兼容")
    add_bullets(
        doc,
        [
            "所有接口使用 /api/v1/ 前缀；",
            "新增字段向后兼容（旧客户端可忽略）；破坏性变更升 v2，并保留 v1 过渡期。",
        ],
    )
    add_h2(doc, "8.5　配置化")
    add_code(
        doc,
        'config.json（学生端与服务端各一份，服务端可远程覆盖）\n'
        '{\n'
        '  "diagnosis": {"total": 15, "k": 6, "t": 5, "s": 2, "p": 2, "weak_threshold": 0.6},\n'
        '  "practice": {"size": 10, "easy": 0.6, "medium": 0.3, "hard": 0.1},\n'
        '  "push_time": "21:30",\n'
        '  "content_version": 12\n'
        "}",
    )
    add_h2(doc, "8.6　发布与部署流程")
    add_code(
        doc,
        "改内容/改代码 → 更新版本号与 CHANGELOG → 跑测试基线\n"
        "  → flutter build apk（App 发布）      → 数据线覆盖安装到 Android 手机\n"
        "  → 服务端升级（备份 SQLite → 拉取代码 → 重启服务）\n"
        "失败回滚：保留上一版 APK 与服务端目录，直接切换回上一版。",
    )
    add_bullets(
        doc,
        [
            "服务端升级脚本 server/deploy.sh：备份 data/ 数据库 → git pull → 重启 uvicorn；",
            "家用电脑开机自启服务；内网穿透（cpolar/花生壳）保证 QQ Webhook 公网可达。",
        ],
    )
    add_h2(doc, "8.7　测试基线（发布前检查清单）")
    add_bullets(
        doc,
        [
            "内容校验脚本通过（章节树/题库一致性）；",
            "规则引擎单元测试通过（诊断/推荐/错题状态机）；",
            "API 回归测试通过（绑定/摘要/状态）；",
            "任一失败则不发版，修复后重跑。",
        ],
    )
    add_h2(doc, "8.8　架构形态：DDD 有界上下文＋模块化单体")
    add_callout(
        doc,
        "架构声明",
        "本项目采用 DDD 领域驱动设计组织代码，但保持模块化单体（单进程部署），现阶段不拆微服务。",
    )
    add_table(
        doc,
        "表12　有界上下文划分",
        ["域", "职责", "主要模块与数据表", "未来切分顺序"],
        [
            ["学习域", "诊断、练习、错题、报告（后续目标/测评）", "features/diagnosis、practice、errorbook、report；diagnosis、answer_record、error_book", "第三切分"],
            ["内容域", "题库、章节树、内容版本", "content/、ContentProvider；question", "第二切分"],
            ["通知域", "绑定、日报、推送", "server/ 通知服务、settings；binding、digest_log、push_log", "第一切分"],
            ["共享支撑", "配置、版本、DB 迁移", "core/、app_config", "不切分"],
        ],
        [1600, 2500, 3260, 2000],
        ["center", "left", "left", "center"],
    )
    add_bullets(
        doc,
        [
            "模块化单体收益：单进程部署简单、升级快，适配家用电脑；微服务在当前单用户规模下只会增加运维成本；",
            "未来演进：若扩展为多学生/多机构，按“通知域→内容域→学习域”顺序以绞杀者模式拆分独立服务，域边界与 /api/v1 版本化接口已就绪；",
            "各域通过防腐层接口（扩展点）交互，拆分时无需改动领域模型。",
        ],
    )

    # ---------------- 九、开发步骤 ----------------
    add_h1(doc, "九、开发步骤（一步一步执行）")
    add_para(
        doc,
        "后续开发严格按以下 7 个阶段推进。每个阶段有明确验收点，通过后再进入下一阶段。"
        "阶段 0 需要你本人完成（安装与账号注册），其余阶段由我们协作完成。",
    )
    add_table(
        doc,
        "表13　开发阶段总览",
        ["阶段", "主要内容", "验收点", "状态", "预计工时"],
        [
            ["阶段0 环境准备", "Flutter/Python 环境、QQ 开放平台机器人", "flutter doctor 通过；机器人创建成功", "✅ 已完成", "0.5–1 天"],
            ["阶段1 工程骨架", "Flutter 工程、目录、路由、主题", "App 可启动并跳转所有页面", "✅ 已完成", "1 天"],
            ["阶段2 数据层", "章节树/题库 JSON、SQLite 建表", "题库导入成功，表可读写", "○ 待开发", "1 天"],
            ["阶段3 诊断报告", "诊断页、评分引擎、报告页", "完成一次诊断并看到雷达图", "○ 待开发", "1–2 天"],
            ["阶段4 练习错题", "练习页、计时、错题状态机", "答错自动入错题本并可重做", "○ 待开发", "2 天"],
            ["阶段5 通知服务", "FastAPI、绑定、摘要、QQ 推送", "家长 QQ 收到测试日报", "○ 待开发", "2 天"],
            ["阶段6 联调部署", "摘要上送、定时推送、部署", "全流程端到端跑通", "○ 待开发", "1 天"],
        ],
        [1000, 2400, 2460, 1300, 2200],
        ["center", "left", "left", "center", "center"],
    )
    add_h2(doc, "9.1　阶段 0：环境准备（需你本人完成）")
    add_numbers(
        doc,
        [
            "安装 Flutter SDK（Windows 版），运行 flutter doctor 确认无阻塞项；",
            "准备 Android 测试手机，开启开发者模式与 USB 调试；",
            "安装 Python 3.10+（开发机已有可复用）；",
            "在 QQ 开放平台注册账号、创建官方机器人，记录 AppID/AppSecret/Token；",
            "准备部署环境：家用电脑（常开）或轻量云服务器，并确认可配置公网回调（或用内网穿透）。",
        ],
    )
    add_para(doc, "状态：✅ 已完成（2026-08-08 验收通过）。实际环境记录：")
    add_bullets(
        doc,
        [
            "Flutter 3.44.9 stable（D:\\flutter）＋Dart 3.12.2；",
            "Android SDK（D:\\Android）：cmdline-tools 13114758、platform-tools 37.0.1、platform android-36、build-tools 36.0.0；Java 21 LTS；全部许可证已接受；",
            "测试手机：iQOO Z5x（V2131A，Android 12，arm64），adb 与 flutter devices 均正常识别；",
            "QQ 官方机器人已注册（AppID/AppSecret/Token 由使用者本地保管，勿外传）；",
            "待办：内网穿透（cpolar/花生壳）注册，阶段 5 需要。",
        ],
    )
    add_h2(doc, "9.2　阶段 1–2：骨架与数据层")
    add_code(
        doc,
        "flutter create math_up_app --org com.example --platforms android   # iOS 后续适配\n"
        "cd math_up_app\n"
        "flutter pub add flutter_math_fork fl_chart sqflite path path_provider http",
    )
    add_bullets(
        doc,
        [
            "按 3.5 目录结构建好 lib/ 与 features/；",
            "编写 chapters.json 与 questions.json（初始建议覆盖必修第一册核心知识点，每知识点 3–5 题，共 50–100 题）；",
            "实现 database.dart 建表（表5）与题库导入。",
        ],
    )
    add_para(doc, "状态：✅ 已完成（2026-08-08，调试 APK 已在 iQOO Z5x 安装运行，6 个占位页跳转验收通过）。实际记录：")
    add_bullets(
        doc,
        [
            "依赖：flutter_math_fork、fl_chart、sqflite、path、path_provider、http、intl（0.20.2，与 flutter_localizations 兼容）；",
            "Android 网络配置：Gradle 腾讯镜像、Maven 阿里云/腾讯镜像、Flutter Maven 中国镜像（storage.flutter-io.cn/download.flutter.io）；",
            "特殊处理：中文路径检查覆盖（android.overridePathCheck=true）、NDK r28c 手动安装（28.2.13676358）、CMake 3.22.1 自动安装；",
            "调试版 APK 约 150MB（含 3 种 ABI）；发布版 arm64 单架构预计 30–40MB。",
        ],
    )
    add_h2(doc, "9.3　阶段 3–4：核心功能")
    add_bullets(
        doc,
        [
            "实现 DiagnosticEngine（规则评分）与报告页（fl_chart 雷达图）；",
            "实现 Recommender（薄弱点＋难度比例）与练习页（计时、即时解析）；",
            "实现错题状态机与周清入口；",
            "验收：完成诊断→报告→练习→错题全流程，纯本地可用。",
        ],
    )
    add_h2(doc, "9.4　阶段 5–6：通知服务与联调")
    add_code(
        doc,
        "pip install fastapi uvicorn apscheduler pydantic requests\n"
        "# 实现 main.py / schemas.py / api/* / application/* / domain/* / infrastructure/channels/qq_bot.py\n"
        "uvicorn app.main:app --host 0.0.0.0 --port 8000",
    )
    add_bullets(
        doc,
        [
            "实现绑定、摘要、状态接口与 QQ Webhook；",
            "实现 APScheduler 每日 21:30 定时推送与失败重试；",
            "联调：学生端配置 Base URL → 生成绑定码 → 家长 QQ 绑定 → 当日摘要上送 → 家长收到日报；",
            "部署：家用电脑开机自启或轻量云 systemd/docker 部署。",
        ],
    )
    add_h2(doc, "9.5　发布与升级流程（持续更新）")
    add_numbers(
        doc,
        [
            "修改内容（content/）或代码后，先更新版本号与 CHANGELOG；",
            "运行测试基线（内容校验＋单元测试＋API 回归），全部通过；",
            "构建并覆盖安装：flutter build apk --release → 数据线连接 Android 手机 → 覆盖安装（禁止卸载）；",
            "服务端升级：备份 SQLite → 拉取代码 → 运行迁移脚本 → 重启服务；",
            "回滚：保留上一版 APK 与服务端目录，出问题时切换回上一版。",
        ],
    )

    # ---------------- 十、测试计划 ----------------
    add_h1(doc, "十、测试计划")
    add_table(
        doc,
        "表14　测试用例计划",
        ["级别", "用例", "预期结果", "执行阶段"],
        [
            ["单元", "诊断评分（含边界：全对/全错/薄弱点判定）", "四因子得分与薄弱点符合规则", "阶段3"],
            ["单元", "推荐选题规则（难度比例、知识点排序）", "10 题按规则选出", "阶段4"],
            ["单元", "错题状态流转（错→重做→观察→掌握）", "状态机正确迁移", "阶段4"],
            ["单元", "摘要生成与离线补传", "队列缓存并自动上送", "阶段5"],
            ["集成", "绑定（正确/错误码/重复绑定/解绑）", "状态与错误码正确", "阶段5"],
            ["集成", "QQ 定时推送与失败重试", "家长收到日报；失败重试 3 次", "阶段5"],
            ["端到端", "首次诊断→练习→错题→家长收日报", "全链路跑通", "阶段6"],
        ],
        [1000, 2500, 3460, 2400],
        ["center", "left", "left", "center"],
    )

    # ---------------- 十一、迭代路线 ----------------
    add_h1(doc, "十一、迭代路线图")
    add_para(
        doc,
        "更新节奏分两条轨道：内容小版本（题库、章节、解析变更）可随时随 App 包发布；"
        "功能按大版本 V1.1–V1.4 演进，每版发布前更新 CHANGELOG 并跑测试基线。",
    )
    add_table(
        doc,
        "表15　后续版本规划",
        ["版本", "新增内容", "对应扩展点"],
        [
            ["V1.1", "SMART 目标、每周任务包、连续打卡激励", "Recommender、SyncService"],
            ["V1.2", "章节测/周测/限时模考（II 卷结构）、失分矩阵", "DiagnosticEngine（P 因子增强）"],
            ["V1.3", "AI 讲解（拆步/类比/图示/反例）、变式生成＋人工审核", "QuestionRepository、AIProvider 新增"],
            ["V1.4", "云同步/账号、教师端预留、日报历史页", "SyncService、TeacherApi"],
        ],
        [1300, 4660, 3400],
        ["center", "left", "left"],
    )

    # ---------------- 十二、风险与预案 ----------------
    add_h1(doc, "十二、风险与预案")
    add_table(
        doc,
        "表16　主要风险与应对",
        ["风险", "影响", "预案"],
        [
            ["QQ 平台收敛主动推送", "定时日报不可用", "家长发“日报”指令主动拉取；邮件通道预案"],
            ["部署设备断电/离线", "日报无法按时推送", "轻量云部署或开机自启；失败次日合并提醒"],
            ["题库质量不足", "诊断与推荐失真", "按章节树逐步补充，教师人工审核后发布"],
            ["未成年人数据合规", "隐私风险", "只上送摘要、本地存储、绑定可解绑"],
            ["公式渲染兼容", "题干显示异常", "flutter_math_fork 统一公式语法，提前适配测试"],
            ["升级破坏本地数据", "学习记录丢失", "强制覆盖安装＋schema 迁移脚本测试"],
        ],
        [1800, 3000, 4560],
        ["center", "left", "left"],
    )

    # ---------------- 附录 ----------------
    add_h1(doc, "附录A　开发协作约定")
    add_bullets(
        doc,
        [
            "每完成一个阶段，把验收结果反馈给我，再进入下一阶段；",
            "代码存放于本工作区（或按阶段 1 创建的 math_up_app 子目录），关键决策记录在本文档对应章节；",
            "遇到与本文档冲突的实际约束，先记录差异再调整，保持架构扩展点不变；",
            "每次发布前运行测试基线（内容校验＋单元测试），并更新 CHANGELOG。",
        ],
    )
    add_h1(doc, "附录B　参考依据")
    add_bullets(
        doc,
        [
            "《高中数学教学能力体系分析报告》（本工作区）；",
            "人教 A 版普通高中教科书·数学（2020 年 5 月第一版）目录；",
            "QQ 开放平台官方机器人文档（以注册时最新说明为准）；",
            "《普通高中数学课程标准》与四川新高考 II 卷考情。",
        ],
    )

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "数学学习提升App开发文档.docx")
    doc.save(out)
    patch_numbering(out)
    print("已生成:", out)
    return out


if __name__ == "__main__":
    build()
